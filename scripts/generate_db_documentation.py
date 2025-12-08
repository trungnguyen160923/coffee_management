#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script để tạo file Word mô tả các bảng trong auth_db và profile_db
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def parse_sql_file(sql_file_path):
    """Parse SQL file và trích xuất thông tin về các bảng"""
    with open(sql_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    tables = []
    table_foreign_keys = {}  # Store FK by table name
    
    # Extract FOREIGN KEYs from ALTER TABLE statements
    alter_table_pattern = r'ALTER TABLE\s+(\w+).*?ADD CONSTRAINT\s+(\w+)\s+FOREIGN\s+KEY\s*\(([^)]+)\)\s+REFERENCES\s+(\w+)\s*\(([^)]+)\)'
    alter_matches = re.finditer(alter_table_pattern, content, re.DOTALL | re.IGNORECASE)
    for match in alter_matches:
        table_name = match.group(1)
        if table_name not in table_foreign_keys:
            table_foreign_keys[table_name] = []
        table_foreign_keys[table_name].append({
            'constraint_name': match.group(2),
            'column': match.group(3).strip(),
            'references_table': match.group(4),
            'references_column': match.group(5).strip()
        })
    
    # Tìm tất cả các CREATE TABLE statements
    create_table_pattern = r'CREATE TABLE\s+(\w+)\s*\((.*?)\)\s*ENGINE'
    matches = re.finditer(create_table_pattern, content, re.DOTALL | re.IGNORECASE)
    
    for match in matches:
        table_name = match.group(1)
        table_body = match.group(2)
        
        # Parse columns - split by comma but handle nested parentheses
        columns = []
        current_col = ""
        paren_depth = 0
        
        for char in table_body:
            if char == '(':
                paren_depth += 1
                current_col += char
            elif char == ')':
                paren_depth -= 1
                current_col += char
            elif char == ',' and paren_depth == 0:
                line = current_col.strip()
                if line:
                    # Process column line
                    if not (line.upper().startswith('KEY') or 
                           line.upper().startswith('CONSTRAINT') or 
                           line.upper().startswith('UNIQUE KEY') or
                           line.upper().startswith('PRIMARY KEY')):
                        col_match = re.match(r'(\w+)\s+(.+)', line, re.DOTALL)
                        if col_match:
                            col_name = col_match.group(1)
                            col_def = col_match.group(2).strip()
                            
                            # Extract data type (handle DECIMAL(12,2), VARCHAR(255), etc.)
                            data_type_match = re.match(r'(\w+(?:\([^)]+\))?)', col_def)
                            data_type = data_type_match.group(1) if data_type_match else col_def.split()[0]
                            
                            # Extract constraints
                            constraints = []
                            col_def_upper = col_def.upper()
                            if 'NOT NULL' in col_def_upper:
                                constraints.append('NOT NULL')
                            if 'AUTO_INCREMENT' in col_def_upper:
                                constraints.append('AUTO_INCREMENT')
                            if 'UNIQUE' in col_def_upper and 'UNIQUE KEY' not in col_def_upper:
                                constraints.append('UNIQUE')
                            if 'DEFAULT' in col_def_upper:
                                default_match = re.search(r'DEFAULT\s+([^,\s]+(?:\([^)]+\))?)', col_def, re.IGNORECASE)
                                if default_match:
                                    constraints.append(f"DEFAULT {default_match.group(1)}")
                            
                            columns.append({
                                'name': col_name,
                                'data_type': data_type,
                                'constraints': constraints,
                                'full_definition': col_def
                            })
                current_col = ""
            else:
                current_col += char
        
        # Handle last column if exists
        if current_col.strip():
            line = current_col.strip()
            if not (line.upper().startswith('KEY') or 
                   line.upper().startswith('CONSTRAINT') or 
                   line.upper().startswith('UNIQUE KEY') or
                   line.upper().startswith('PRIMARY KEY')):
                col_match = re.match(r'(\w+)\s+(.+)', line, re.DOTALL)
                if col_match:
                    col_name = col_match.group(1)
                    col_def = col_match.group(2).strip()
                    data_type_match = re.match(r'(\w+(?:\([^)]+\))?)', col_def)
                    data_type = data_type_match.group(1) if data_type_match else col_def.split()[0]
                    constraints = []
                    col_def_upper = col_def.upper()
                    if 'NOT NULL' in col_def_upper:
                        constraints.append('NOT NULL')
                    if 'AUTO_INCREMENT' in col_def_upper:
                        constraints.append('AUTO_INCREMENT')
                    if 'UNIQUE' in col_def_upper and 'UNIQUE KEY' not in col_def_upper:
                        constraints.append('UNIQUE')
                    if 'DEFAULT' in col_def_upper:
                        default_match = re.search(r'DEFAULT\s+([^,\s]+(?:\([^)]+\))?)', col_def, re.IGNORECASE)
                        if default_match:
                            constraints.append(f"DEFAULT {default_match.group(1)}")
                    columns.append({
                        'name': col_name,
                        'data_type': data_type,
                        'constraints': constraints,
                        'full_definition': col_def
                    })
        
        # Extract PRIMARY KEY from table body
        primary_key_match = re.search(r'PRIMARY KEY\s*\(([^)]+)\)', table_body, re.IGNORECASE)
        primary_keys = []
        if primary_key_match:
            primary_keys = [pk.strip() for pk in primary_key_match.group(1).split(',')]
        
        # Extract FOREIGN KEYs from table body
        foreign_keys = []
        fk_pattern = r'CONSTRAINT\s+(\w+)\s+FOREIGN\s+KEY\s*\(([^)]+)\)\s+REFERENCES\s+(\w+)\s*\(([^)]+)\)'
        fk_matches = re.finditer(fk_pattern, table_body, re.IGNORECASE)
        for fk_match in fk_matches:
            foreign_keys.append({
                'constraint_name': fk_match.group(1),
                'column': fk_match.group(2).strip(),
                'references_table': fk_match.group(3),
                'references_column': fk_match.group(4).strip()
            })
        
        # Add FOREIGN KEYs from ALTER TABLE
        if table_name in table_foreign_keys:
            foreign_keys.extend(table_foreign_keys[table_name])
        
        # Extract INDEXes
        indexes = []
        # Match KEY idx_name (columns)
        index_pattern = r'KEY\s+(\w+)\s*\(([^)]+)\)'
        index_matches = re.finditer(index_pattern, table_body, re.IGNORECASE)
        for idx_match in index_matches:
            indexes.append({
                'name': idx_match.group(1),
                'columns': [col.strip() for col in idx_match.group(2).split(',')]
            })
        
        # Match UNIQUE KEY
        unique_key_pattern = r'UNIQUE KEY\s+(\w+)\s*\(([^)]+)\)'
        unique_matches = re.finditer(unique_key_pattern, table_body, re.IGNORECASE)
        for idx_match in unique_matches:
            indexes.append({
                'name': idx_match.group(1),
                'columns': [col.strip() for col in idx_match.group(2).split(',')]
            })
        
        tables.append({
            'name': table_name,
            'columns': columns,
            'primary_keys': primary_keys,
            'foreign_keys': foreign_keys,
            'indexes': indexes
        })
    
    return tables

def add_table_heading(doc, text, level=1):
    """Thêm heading cho document"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.name = 'Times New Roman'
    return heading

def add_table_to_doc(doc, table_info):
    """Thêm thông tin bảng vào document"""
    # Tên bảng
    doc.add_heading(table_info['name'], level=2)
    
    # Mô tả các cột
    doc.add_heading('Các cột (Columns)', level=3)
    
    # Tạo bảng cho columns
    cols_table = doc.add_table(rows=1, cols=4)
    cols_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = cols_table.rows[0].cells
    header_cells[0].text = 'Tên cột'
    header_cells[1].text = 'Kiểu dữ liệu'
    header_cells[2].text = 'Ràng buộc'
    header_cells[3].text = 'Ghi chú'
    
    # Format header
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Add column data
    for col in table_info['columns']:
        row_cells = cols_table.add_row().cells
        row_cells[0].text = col['name']
        row_cells[1].text = col['data_type']
        row_cells[2].text = ', '.join(col['constraints']) if col['constraints'] else '-'
        row_cells[3].text = ''
    
    # Primary Keys
    if table_info['primary_keys']:
        doc.add_heading('Khóa chính (Primary Keys)', level=3)
        p = doc.add_paragraph(', '.join(table_info['primary_keys']))
        p.style = 'List Bullet'
    
    # Foreign Keys
    if table_info['foreign_keys']:
        doc.add_heading('Khóa ngoại (Foreign Keys)', level=3)
        fk_table = doc.add_table(rows=1, cols=4)
        fk_table.style = 'Light Grid Accent 1'
        
        fk_header = fk_table.rows[0].cells
        fk_header[0].text = 'Tên ràng buộc'
        fk_header[1].text = 'Cột'
        fk_header[2].text = 'Tham chiếu bảng'
        fk_header[3].text = 'Tham chiếu cột'
        
        for cell in fk_header:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        for fk in table_info['foreign_keys']:
            fk_row = fk_table.add_row().cells
            fk_row[0].text = fk['constraint_name']
            fk_row[1].text = fk['column']
            fk_row[2].text = fk['references_table']
            fk_row[3].text = fk['references_column']
    
    # Indexes
    if table_info['indexes']:
        doc.add_heading('Chỉ mục (Indexes)', level=3)
        idx_table = doc.add_table(rows=1, cols=2)
        idx_table.style = 'Light Grid Accent 1'
        
        idx_header = idx_table.rows[0].cells
        idx_header[0].text = 'Tên chỉ mục'
        idx_header[1].text = 'Cột'
        
        for cell in idx_header:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        for idx in table_info['indexes']:
            idx_row = idx_table.add_row().cells
            idx_row[0].text = idx['name']
            idx_row[1].text = ', '.join(idx['columns'])
    
    doc.add_paragraph()  # Add spacing

def generate_word_document():
    """Tạo file Word từ các SQL files"""
    # Đường dẫn đến các file SQL
    base_path = Path(__file__).parent.parent
    auth_db_path = base_path / 'sql' / 'auth_db.sql'
    profile_db_path = base_path / 'sql' / 'profile_db.sql'
    
    if not auth_db_path.exists():
        print(f"Lỗi: Không tìm thấy file {auth_db_path}")
        return
    
    if not profile_db_path.exists():
        print(f"Lỗi: Không tìm thấy file {profile_db_path}")
        return
    
    # Parse SQL files
    print("Đang phân tích auth_db.sql...")
    auth_tables = parse_sql_file(auth_db_path)
    
    print("Đang phân tích profile_db.sql...")
    profile_tables = parse_sql_file(profile_db_path)
    
    # Tạo document
    doc = Document()
    
    # Set font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('TÀI LIỆU MÔ TẢ CƠ SỞ DỮ LIỆU', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Auth Database và Profile Database')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    
    # Auth Database section
    doc.add_page_break()
    add_table_heading(doc, 'AUTH DATABASE', 1)
    doc.add_paragraph(f'Tổng số bảng: {len(auth_tables)}')
    doc.add_paragraph()
    
    for table in auth_tables:
        add_table_to_doc(doc, table)
    
    # Profile Database section
    doc.add_page_break()
    add_table_heading(doc, 'PROFILE DATABASE', 1)
    doc.add_paragraph(f'Tổng số bảng: {len(profile_tables)}')
    doc.add_paragraph()
    
    for table in profile_tables:
        add_table_to_doc(doc, table)
    
    # Save document
    output_path = base_path / 'docs' / 'database_documentation.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    
    print(f"\n✓ Đã tạo file Word thành công tại: {output_path}")
    print(f"  - Auth Database: {len(auth_tables)} bảng")
    print(f"  - Profile Database: {len(profile_tables)} bảng")

if __name__ == '__main__':
    try:
        generate_word_document()
    except ImportError as e:
        print("Lỗi: Thiếu thư viện cần thiết.")
        print("Vui lòng cài đặt: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

